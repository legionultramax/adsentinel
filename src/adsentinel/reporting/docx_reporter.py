"""
ADSentinel DOCX Report Generator (Final Version)
Supports two modes:
- executive: Clean, management-friendly summary with MITRE mapping
- full: Detailed technical report with all findings

Usage:
    from adsentinel.reporting.docx_reporter import generate_docx_report
    generate_docx_report(scan_result, "report.docx", mode="executive")
"""

from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Any, List

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from adsentinel import __version__
from adsentinel.logging_config import get_logger
from adsentinel.models.severity import Severity

logger = get_logger(__name__)


def set_cell_shading(cell, fill_color: str):
    """Set background color for a table cell."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill_color)
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_colored_heading(doc, text: str, level: int = 1, color: str = "1F4E79"):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor.from_string(color)
    return heading


def generate_docx_report(scan_result: Any, output_path: str, mode: str = "executive") -> None:
    """
    Generate DOCX report in executive or full mode.

    Args:
        scan_result: ADSentinel scan result object
        output_path: Path to save .docx file
        mode: "executive" or "full"
    """
    if mode not in ("executive", "full"):
        raise ValueError("mode must be 'executive' or 'full'")

    doc = Document()

    # Set comfortable margins
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.6)
        section.right_margin = Cm(1.6)

    findings: List[Any] = getattr(scan_result, "all_findings", [])
    score: float = getattr(scan_result, "posture_score", 0.0)
    grade: str = getattr(scan_result, "grade", "F")
    domain: str = "Unknown"
    if hasattr(scan_result, "config_summary"):
        domain = scan_result.config_summary.get("domain", "Unknown")

    if mode == "executive":
        _build_executive_report(doc, scan_result, findings, score, grade, domain)
    else:
        _build_full_report(doc, scan_result, findings, score, grade, domain)

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    logger.info("docx_report_generated", path=output_path, mode=mode, findings=len(findings))


def _build_executive_report(doc, scan_result, findings, score, grade, domain):
    """Clean executive summary report."""

    # Title
    title = doc.add_heading("Active Directory Security Assessment", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        run.font.size = Pt(22)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Executive Summary with MITRE ATT&CK Mapping")
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Date: {datetime.now().strftime('%d %b %Y')}  |  Domain: {domain}").italic = True

    doc.add_paragraph()

    # Risk Rating
    add_colored_heading(doc, "Overall Risk Rating", level=1, color="C00000")

    risk_table = doc.add_table(rows=4, cols=2)
    risk_table.style = "Table Grid"

    data = [
        ("Posture Score", f"{score:.1f}/100   (Grade: {grade})"),
        ("Risk Level", "CRITICAL" if score < 50 else "HIGH"),
        ("Total Findings", str(len(findings))),
        ("Breakdown", f"Critical: {getattr(scan_result, 'critical_count', 0)}  |  "
                      f"High: {getattr(scan_result, 'high_count', 0)}  |  "
                      f"Medium: {getattr(scan_result, 'medium_count', 0)}"),
    ]

    for i, (label, value) in enumerate(data):
        risk_table.cell(i, 0).text = label
        risk_table.cell(i, 1).text = value
        for p in risk_table.cell(i, 0).paragraphs:
            for run in p.runs:
                run.bold = True

    set_cell_shading(risk_table.cell(1, 1), "FFCCCC")

    doc.add_paragraph()

    # MITRE ATT&CK Mapping
    add_colored_heading(doc, "MITRE ATT&CK Mapping – Key Techniques", level=1)

    mitre_table = doc.add_table(rows=1, cols=4)
    mitre_table.style = "Table Grid"

    headers = ["Attack Path", "MITRE ID", "Technique Name", "Severity"]
    for i, h in enumerate(headers):
        mitre_table.rows[0].cells[i].text = h
        set_cell_shading(mitre_table.rows[0].cells[i], "1F4E79")
        for p in mitre_table.rows[0].cells[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(9)

    mitre_data = [
        ("AD CS ESC1 – Certificate as Domain Admin", "T1649", "Steal or Forge Authentication Certificates", "CRITICAL"),
        ("", "T1558", "Steal or Forge Kerberos Tickets", "CRITICAL"),
        ("Coercion + ESC8 + DC Certificate Relay", "T1557", "Adversary-in-the-Middle (NTLM Relay)", "CRITICAL"),
        ("", "T1187", "Forced Authentication", "HIGH"),
        ("RBCD Abuse via Machine Account Quota", "T1134.002", "Create Process with Token", "HIGH"),
    ]

    for attack, tid, name, sev in mitre_data:
        row = mitre_table.add_row()
        row.cells[0].text = attack
        row.cells[1].text = tid
        row.cells[2].text = name
        row.cells[3].text = sev
        if sev == "CRITICAL":
            set_cell_shading(row.cells[3], "FFCCCC")
        elif sev == "HIGH":
            set_cell_shading(row.cells[3], "FFE699")

    doc.add_paragraph()

    # Top Risks
    add_colored_heading(doc, "Top Risks & Business Impact", level=1)

    risks = [
        ("AD CS ESC1 (Critical)", 
         "Any authenticated user can request a certificate for a Domain Admin account. Leads to full domain compromise via Kerberos PKINIT."),
        ("Coercion + NTLM Relay to AD CS (Critical)",
         "Attacker can coerce authentication and relay it to obtain a Domain Controller certificate, enabling DCSync and complete domain takeover."),
        ("RBCD Abuse via Machine Account Quota (High)",
         "Regular domain users can create computer accounts and abuse delegation to impersonate high-privileged accounts."),
    ]

    for title, impact in risks:
        p = doc.add_paragraph()
        p.add_run("• " + title + ": ").bold = True
        p.add_run(impact)

    # Recommended Actions
    add_colored_heading(doc, "Recommended Immediate Actions", level=1)

    p = doc.add_paragraph()
    p.add_run("Within 7 Days:").bold = True
    doc.add_paragraph("Harden AD CS certificate templates – remove dangerous permissions (ESC1).", style="List Bullet")
    doc.add_paragraph("Block NTLM relay to AD CS web enrollment endpoints.", style="List Bullet")

    p = doc.add_paragraph()
    p.add_run("Within 30 Days:").bold = True
    doc.add_paragraph("Restrict who can create machine accounts in the domain.", style="List Bullet")
    doc.add_paragraph("Implement monitoring for suspicious certificate requests and delegation changes.", style="List Bullet")

    # Conclusion
    add_colored_heading(doc, "Conclusion", level=1)
    p = doc.add_paragraph()
    p.add_run("The Active Directory security posture is currently ")
    run = p.add_run("Critical")
    run.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    p.add_run(". Multiple MITRE ATT&CK techniques can be chained to achieve full domain compromise. Immediate remediation of AD Certificate Services misconfigurations is strongly recommended.")

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.add_run(f"Generated by ADSentinel v{__version__}  |  {datetime.now().strftime('%Y-%m-%d %H:%M')}").italic = True


def _build_full_report(doc, scan_result, findings, score, grade, domain):
    """Detailed report with all findings."""

    _build_executive_report(doc, scan_result, findings, score, grade, domain)

    doc.add_page_break()
    add_colored_heading(doc, "Detailed Findings", level=1, color="1F4E79")

    if not findings:
        doc.add_paragraph("No findings were identified.")
        return

    # Severity summary table
    sev_counts = {"CRITICAL": 0, "HIGH": 0, "MEDIUM": 0, "LOW": 0, "INFO": 0}
    for f in findings:
        sev = getattr(f, "severity", Severity.INFO).value
        if sev in sev_counts:
            sev_counts[sev] += 1

    summary_table = doc.add_table(rows=2, cols=5)
    summary_table.style = "Table Grid"
    severities = ["CRITICAL", "HIGH", "MEDIUM", "LOW", "INFO"]
    colors = ["FFCCCC", "FFE699", "FFFFCC", "CCE5FF", "E0E0E0"]

    for i, sev in enumerate(severities):
        summary_table.cell(0, i).text = sev
        summary_table.cell(1, i).text = str(sev_counts[sev])
        set_cell_shading(summary_table.cell(0, i), colors[i])
        set_cell_shading(summary_table.cell(1, i), colors[i])
        for p in summary_table.cell(0, i).paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)

    doc.add_paragraph()

    # Detailed findings
    severity_order = {"CRITICAL": 0, "HIGH": 1, "MEDIUM": 2, "LOW": 3, "INFO": 4}
    sorted_findings = sorted(findings, key=lambda f: severity_order.get(getattr(f, "severity", Severity.INFO).value, 5))

    for f in sorted_findings:
        sev = getattr(f, "severity", Severity.INFO).value
        fid = getattr(f, "id", "")
        title = getattr(f, "title", "Untitled")
        desc = getattr(f, "description", "")
        affected = getattr(f, "affected_objects", [])
        remediation = getattr(f, "remediation", None)
        compliance = getattr(f, "compliance", None)

        p = doc.add_paragraph()
        run = p.add_run(f"[{sev}] {fid} – {title}")
        run.bold = True
        if sev == "CRITICAL":
            run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        elif sev == "HIGH":
            run.font.color.rgb = RGBColor(0xED, 0x7D, 0x31)

        if desc:
            doc.add_paragraph(desc)

        if affected:
            p = doc.add_paragraph()
            p.add_run("Affected Objects: ").bold = True
            p.add_run(f"{len(affected)} object(s)")
            for obj in affected[:8]:
                name = getattr(obj, "sam_account_name", None) or getattr(obj, "dn", "Unknown")
                doc.add_paragraph(f"   • {name}", style="List Bullet")
            if len(affected) > 8:
                doc.add_paragraph(f"   ... and {len(affected) - 8} more", style="List Bullet")

        if remediation and getattr(remediation, "description", ""):
            p = doc.add_paragraph()
            p.add_run("Remediation: ").bold = True
            p.add_run(remediation.description)
            if getattr(remediation, "powershell_command", ""):
                pre = doc.add_paragraph()
                pre.add_run(remediation.powershell_command).font.name = "Consolas"
                pre.paragraph_format.left_indent = Cm(0.5)

        if compliance:
            mitre = getattr(compliance, "mitre_attack", [])
            if mitre:
                tags = ", ".join([m.technique_id for m in mitre])
                p = doc.add_paragraph()
                p.add_run("MITRE ATT&CK: ").bold = True
                p.add_run(tags)

        doc.add_paragraph()
